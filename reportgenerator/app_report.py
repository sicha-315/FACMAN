import eventlet
eventlet.monkey_patch()

import os
from flask import Flask, render_template, request, jsonify, send_file, make_response
from flask_socketio import SocketIO
from influxdb_client import InfluxDBClient
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
from io import BytesIO
from docx.shared import Inches
import base64
import openai
from collections import defaultdict
from datetime import datetime, timezone, timedelta
import json
import traceback
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook

matplotlib.use("Agg") 

# 한글 폰트 설정
font_path = "C:/Windows/Fonts/malgun.ttf"  # Windows 기본 한글 폰트 (맑은 고딕)
font_prop = fm.FontProperties(fname=font_path, size=12)
plt.rc('font', family=font_prop.get_name())

# ===============================
# 환경 변수 및 클라이언트 설정
# ===============================
load_dotenv()

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="eventlet")

INFLUX_URL = os.getenv("INFLUX_URL")
INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
INFLUX_ORG = os.getenv("INFLUX_ORG")
influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)
openai.api_key = os.getenv("OPENAI_API_KEY")

# ===============================
# 한글 기간 문자열 변환 함수
# ===============================
def normalize_range(range_str):
    kor_to_influx = {
        "1시간": "1h", "3시간": "3h", "6시간": "6h", "9시간": "9h",
        "1일": "1d", "7일": "7d", "31일": "31d"
    }
    return kor_to_influx.get(range_str, range_str)

# ===============================
# 실시간 상태 조회 및 전송
# ===============================
def get_recent_status(bucket):
    query = f'''
    from(bucket: "{bucket}")
      |> range(start: -30s)
      |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type")
      |> sort(columns: ["_time"], desc: true)
      |> limit(n: 3)
    '''
    result = influx_client.query_api().query(org=INFLUX_ORG, query=query)
    events = [record.get_value() for table in result for record in table.records]
    return events if events else None

def emit_status():
    prev_events = {"P1-A": None, "P1-B": None, "P2-A": None, "P2-B": None}
    while True:
        for key in prev_events.keys():
            events = get_recent_status(f"{key}_status")
            if events:
                latest = events[0]
                if latest != prev_events[key]:
                    socketio.emit('status_update', {key: {'event_type': latest}})
                    prev_events[key] = latest
        socketio.sleep(1)

@socketio.on('connect')
def handle_connect():
    for key in ["P1", "P2"]:
        events = get_recent_status(f"{key}_status")
        if events:
            latest = events[0]
            socketio.emit('status_update', {f"{key}-A": {'event_type': latest}})

# ===============================
# 라우팅
# ===============================
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/report")
def report_page():
    return render_template("report.html")

# ===============================
# 보고서 생성 API (다중 공정 대응)
# ===============================
@app.route("/generate_report", methods=["POST"])
def generate_report():
    try:
        data = request.get_json()
        processes = data.get("processes", [])
        range_str = normalize_range(data.get("range"))

        all_reports = []
        for process in processes:
            if "/" in range_str:
                start, end = range_str.split("/")
                range_clause = f'|> range(start: time(v: "{start}"), stop: time(v: "{end}"))'
                start_time = datetime.fromisoformat(start.replace("+09:00", ""))
                end_time = datetime.fromisoformat(end.replace("+09:00", ""))
            else:
                range_clause = f'|> range(start: -{range_str})'
                start_time, end_time = None, None

            query = f'''
            from(bucket: "{process}_status")
              {range_clause}
              |> filter(fn: (r) => r._measurement == "status_log")
              |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
              |> keep(columns: ["_time", "available", "event_type"])
            '''
            tables = influx_client.query_api().query(query)

            total, available_sum, failure_count = 0, 0, 0
            time_labels, available_values, failure_values = [], [], []
            failure_hourly = defaultdict(int)

            KST = timezone(timedelta(hours=9))

            for table in tables:
                for record in table.records:
                    record_time = record.values["_time"].astimezone(KST).replace(tzinfo=None)
                    if start_time and end_time:
                        if not (start_time <= record_time <= end_time):
                            continue

                    available = record.values.get("available", 0)
                    event_type = record.values.get("event_type", "")
                    timestamp = record_time.strftime("%H:%M")
                    hour_label = record_time.strftime("%H시대")

                    total += 1
                    available_sum += available
                    if event_type == "failure":
                        failure_count += 1
                        failure_hourly[hour_label] += 1

                    time_labels.append(timestamp)
                    available_values.append(round(available, 2))
                    failure_values.append(1 if event_type == "failure" else 0)

            avg_avail = round((available_sum / total) * 100, 1) if total else 0
            failure_table_labels = list(failure_hourly.keys())
            failure_table_counts = list(failure_hourly.values())

            # ✅ 생산실적 계산
            p0_query = f'''
            from(bucket: "process")
              {range_clause}
              |> filter(fn: (r) => r._measurement == "process_log" and r.process_id == "P0")
              |> group()
              |> count()
            '''
            p3_query = f'''
            from(bucket: "process")
              {range_clause}
              |> filter(fn: (r) => r._measurement == "process_log" and r.process_id == "P3")
              |> group()
              |> count()
            '''
            p0_count = sum([record.get_value() for table in influx_client.query_api().query(org=INFLUX_ORG, query=p0_query) for record in table.records])
            p3_count = sum([record.get_value() for table in influx_client.query_api().query(org=INFLUX_ORG, query=p3_query) for record in table.records])
            production_rate = round((p3_count / p0_count) * 100, 1) if p0_count else 0

            prompt = f"""
공정명: {process}
기간: 최근 {range_str}
가동률 평균: {avg_avail}%
고장 횟수: {failure_count}회
생산실적: 투입 {p0_count}개 → 산출 {p3_count}개 (양품률 {production_rate}%)

위 데이터를 바탕으로 제조 공정 보고서를 작성해줘.
아래 각 항목에 대해 글머리 기호 '-'로 중요한 내용을 포함하도록 작성하고, '보고서 작성자:', '이상입니다' 등의 표현은 절대 포함하지 마. 
항목:
1. 공정 요약
2. 주요 이슈
3. 대응 조치
4. 향후 제언
"""
            response = openai.ChatCompletion.create(
                model="gpt-4-1106-preview",
                messages=[{"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                          {"role": "user", "content": prompt}]
            )

            summary_resp = openai.ChatCompletion.create(
                model="gpt-4-1106-preview",
                messages=[
                    {"role": "system", "content": "너는 간결한 제조 보고서 요약가야."},
                    {"role": "user", "content": prompt}
                ]
            )
            full_report = response.choices[0].message.content.strip()
            
            # ✅ MTBF 데이터 수집
            mtbf_resp = app.test_client().post("/get_mtbf_data", json={"process": process, "range": range_str})
            mtbf_json = mtbf_resp.get_json() if mtbf_resp.status_code == 200 else {}

            # ✅ MTTR 데이터 수집
            mttr_resp = app.test_client().post("/get_mttr_data", json={"process": process, "range": range_str})
            mttr_json = mttr_resp.get_json() if mttr_resp.status_code == 200 else {}

            all_reports.append({
                "process": process,
                "summary": full_report,
                "report": response.choices[0].message.content,
                "labels": time_labels,
                "available": available_values,
                "failures": failure_values,
                "failureLabels": failure_table_labels,
                "failureCounts": failure_table_counts,
                "production": {
                    "input": p0_count,
                    "output": p3_count,
                    "rate": production_rate
                },
                **mtbf_json,
                **mttr_json
            })


        return jsonify({"reports": all_reports})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ===============================
# 생산실적 계산 API
# ===============================
@app.route("/get_production_data", methods=["POST"])
def get_production_data():
    try:
        data = request.get_json()
        range_str = data.get("range")

        if not range_str:
            return jsonify({"error": "Missing range parameter"}), 400

        # 시간 범위 파싱
        if "/" in range_str:
            start_str, end_str = range_str.split("/")
            start = datetime.fromisoformat(start_str.replace("Z", "").replace("+09:00", ""))
            end = datetime.fromisoformat(end_str.replace("Z", "").replace("+09:00", ""))
            range_clause = f'|> range(start: time(v: "{start_str}"), stop: time(v: "{end_str}"))'
        else:
            range_clause = f'|> range(start: -{range_str})'

        def count_query(process_id):
            query = f'''
            from(bucket: "process")
              {range_clause}
              |> filter(fn: (r) => r._measurement == "process_log" and r.process_id == "{process_id}")
              |> group()
              |> count()
            '''
            result = influx_client.query_api().query(org=INFLUX_ORG, query=query)
            return sum([record.get_value() for table in result for record in table.records])

        p0_count = count_query("P0")
        p3_count = count_query("P3")

        ratio = round((p3_count / p0_count) * 100, 1) if p0_count else 0

        return jsonify({
            "input": p0_count,
            "output": p3_count,
            "rate": ratio
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ===============================
# 다운타임 계산 API
# ===============================
@app.route("/get_downtime_data", methods=["POST"])
def get_downtime_data():
    data = request.json
    process = data.get("process")
    range_str = data.get("range")

    if not process or not range_str:
        return jsonify({"error": "Missing process or range"}), 400

    # 시간 범위 파싱
    if "/" in range_str:
        start_str, end_str = range_str.split("/")
        start = datetime.fromisoformat(start_str.replace("Z", "").replace("+09:00", ""))
        end = datetime.fromisoformat(end_str.replace("Z", "").replace("+09:00", ""))
        range_clause = f'|> range(start: time(v: "{start_str}"), stop: time(v: "{end_str}"))'
    else:
        range_clause = f'|> range(start: -{range_str})'
        now = datetime.utcnow()
        if "h" in range_str:
            hours = int(range_str.replace("h", ""))
            start = now - timedelta(hours=hours)
        elif "d" in range_str:
            days = int(range_str.replace("d", ""))
            start = now - timedelta(days=days)
        end = now

    downtime_events = ["failure", "repair"]

    query = f'''
        from(bucket: "{process}_status")
        {range_clause}
        |> filter(fn: (r) => r._measurement == "status_log")
        |> filter(fn: (r) => r._field == "available")
        |> sort(columns: ["_time"])
        |> difference(columns: ["_value"], keepFirst: true)
        |> filter(fn: (r) => r._value != 0)
        |> map(fn: (r) => ({{ r with time_ns: uint(v: r._time) }}))
        |> difference(columns: ["time_ns"])
        |> map(fn: (r) => ({{ r with time_diff_sec: float(v: r.time_ns) / 1000000000.0 }}))
        |> filter(fn: (r) => r._value == 1)
        |> sum(column: "time_diff_sec")
        '''

    tables = influx_client.query_api().query(org=INFLUX_ORG, query=query)

    failure_total = 0
    repair_total = 0
    failure_by_hour = defaultdict(float)
    repair_by_hour = defaultdict(float)

    current_event = None
    current_start = None
    KST = timezone(timedelta(hours=9))

    for table in tables:
        for record in table.records:
            event_type = record.get_value()
            timestamp = record.get_time().astimezone(KST).replace(tzinfo=None)

            if event_type in downtime_events:
                current_event = event_type
                current_start = timestamp
            elif event_type == "processing" and current_event and current_start:
                diff = (timestamp - current_start).total_seconds() / 60
                hour_label = current_start.strftime("%H시")

                if current_event == "failure":
                    failure_total += diff
                    failure_by_hour[hour_label] += diff
                elif current_event == "repair":
                    repair_total += diff
                    repair_by_hour[hour_label] += diff

                current_event = None
                current_start = None

    return jsonify({
        "failure_total": round(failure_total, 1),
        "repair_total": round(repair_total, 1),
        "hourly_labels": sorted(set(list(failure_by_hour.keys()) + list(repair_by_hour.keys()))),
        "failure_by_hour": [round(failure_by_hour.get(h, 0), 1) for h in sorted(set(list(failure_by_hour.keys()) + list(repair_by_hour.keys())))],
        "repair_by_hour": [round(repair_by_hour.get(h, 0), 1) for h in sorted(set(list(failure_by_hour.keys()) + list(repair_by_hour.keys())))]
    })

# ===============================
# MTBF 계산 API
# ===============================
@app.route("/get_mtbf_data", methods=["POST"])
def get_mtbf_data():
    try:
        data = request.get_json()
        process = data.get("process")
        range_str = data.get("range")

        if not process or not range_str:
            return jsonify({"error": "Missing process or range"}), 400

        # 시간 범위 파싱
        if "/" in range_str:
            start_str, end_str = range_str.split("/")
            start = datetime.fromisoformat(start_str.replace("Z", "").replace("+09:00", ""))
            end = datetime.fromisoformat(end_str.replace("Z", "").replace("+09:00", ""))
            range_clause = f'|> range(start: time(v: "{start_str}"), stop: time(v: "{end_str}"))'
        else:
            range_clause = f'|> range(start: -{range_str})'
            now = datetime.utcnow()
            if "h" in range_str:
                hours = int(range_str.replace("h", ""))
                start = now - timedelta(hours=hours)
            elif "d" in range_str:
                days = int(range_str.replace("d", ""))
                start = now - timedelta(days=days)
            end = now

        query = f'''
        from(bucket: "{process}_status")
          {range_clause}
          |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type")
          |> sort(columns: ["_time"])
        '''

        tables = influx_client.query_api().query(org=INFLUX_ORG, query=query)

        failure_count = 0
        total_processing_minutes = 0
        current_event = None
        current_start = None

        KST = timezone(timedelta(hours=9))
        failure_states = ["failure"]  # ✅ 유지보수 제외

        for table in tables:
            for record in table.records:
                event_type = record.get_value()
                timestamp = record.get_time().astimezone(KST).replace(tzinfo=None)

                if event_type == "processing":
                    current_event = "processing"
                    current_start = timestamp
                elif event_type in failure_states and current_event == "processing" and current_start:
                    diff = (timestamp - current_start).total_seconds() / 60
                    total_processing_minutes += diff
                    failure_count += 1
                    current_event = None
                    current_start = None

        mtbf = round(total_processing_minutes / failure_count, 1) if failure_count else 0

        return jsonify({
            "failure_count": failure_count,
            "total_processing_minutes": round(total_processing_minutes, 1),
            "mtbf_minutes": mtbf,
            "summary_text": f"""🔁 MTBF 요약
---------------------------
고장 횟수: {failure_count}회
총 운영 시간: {round(total_processing_minutes, 1)}분
고장 간 평균 시간 (MTBF): {mtbf}분"""
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ===============================
# MTTR 계산 API
# ===============================
@app.route("/get_mttr_data", methods=["POST"])
def get_mttr_data():
    try:
        data = request.get_json()
        process = data.get("process")
        range_str = data.get("range")

        if not process or not range_str:
            return jsonify({"error": "Missing process or range"}), 400

        if "/" in range_str:
            start_str, end_str = range_str.split("/")
            range_clause = f'|> range(start: time(v: "{start_str}"), stop: time(v: "{end_str}"))'
        else:
            range_clause = f'|> range(start: -{range_str})'

        query = f'''
        from(bucket: "{process}_status")
          {range_clause}
          |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_status" and r._value == "start")
        '''

        finish_query = f'''
        from(bucket: "{process}_status")
          {range_clause}
          |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_status" and r._value == "finish")
        '''

        start_times = [
            record.get_time().astimezone(timezone(timedelta(hours=9))).replace(tzinfo=None)
            for table in influx_client.query_api().query(org=INFLUX_ORG, query=query)
            for record in table.records
        ]

        finish_times = [
            record.get_time().astimezone(timezone(timedelta(hours=9))).replace(tzinfo=None)
            for table in influx_client.query_api().query(org=INFLUX_ORG, query=finish_query)
            for record in table.records
        ]

        # ✅ 짝수 맞춰 계산
        min_len = min(len(start_times), len(finish_times))
        repair_durations = [
            (finish_times[i] - start_times[i]).total_seconds() / 60
            for i in range(min_len)
            if finish_times[i] > start_times[i]
        ]

        total_repair_time = sum(repair_durations)
        repair_count = len(repair_durations)
        mttr = round(total_repair_time / repair_count, 1) if repair_count else 0

        return jsonify({
            "repair_count": repair_count,
            "total_repair_minutes": round(total_repair_time, 1),
            "mttr_minutes": mttr,
            "summary_text": f"""🔧 MTTR 요약
---------------------------
수리 횟수: {repair_count}회
총 수리 시간: {round(total_repair_time, 1)}분
평균 수리 시간 (MTTR): {mttr}분"""
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ===============================
# DOCX 다운로드 API 
# ===============================
# ==============================
# Chart Generator Functions
# ==============================
from io import BytesIO
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import request, send_file, jsonify
import json

# ✅ Downtime 기준에서 maintenance 제외

def create_donut_chart(percent, labels, colors):
    fig, ax = plt.subplots(figsize=(2.5, 2.5), dpi=100)
    ax.pie([percent, 100 - percent], labels=labels, colors=colors,
           startangle=90, wedgeprops={"width": 0.4})
    ax.set(aspect="equal")
    buf = BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

def create_production_bar(input_cnt, output_cnt):
    fig, ax = plt.subplots(figsize=(3, 2.5), dpi=100)
    ax.bar(["투입량", "산출량"], [input_cnt, output_cnt], color=["blue", "green"])
    ax.set_title("생산실적")
    buf = BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

def create_failure_line_chart(labels, values):
    fig, ax = plt.subplots(figsize=(4, 3), dpi=100)
    ax.plot(labels, values, marker='o', color='red')
    ax.set_title("고장 발생 분포")
    ax.set_xlabel("시간대")
    ax.set_ylabel("건수")
    ax.grid(True)
    plt.xticks(rotation=45, fontsize=8)
    buf = BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

def create_total_downtime_pie(failure_min, repair_min, operating_min):
    labels = ['고장 시간', '수리 시간', '운영 시간']
    values = [failure_min, repair_min, operating_min]
    colors = ['#ff6b6b', '#ffa94d', '#8ce99a']
    fig, ax = plt.subplots(figsize=(3, 3), dpi=100)
    ax.pie(values, labels=labels, colors=colors, startangle=90, autopct='%1.1f%%')
    ax.set(aspect="equal")
    buf = BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

def create_downtime_bar_chart(labels, failureData):
    fig, ax = plt.subplots(figsize=(5, 3), dpi=100)
    ax.bar(labels, failureData, color="red")
    ax.set_title("시간대별 고장 다운타임")
    ax.set_xlabel("시간대")
    ax.set_ylabel("다운타임 (분)")
    ax.grid(True)
    plt.xticks(rotation=45, fontsize=8)
    buf = BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf

@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        text = request.form.get("report", "")
        failure_labels = json.loads(request.form.get("failureLabels", "[]"))
        failure_counts = json.loads(request.form.get("failureCounts", "[]"))
        report_data = json.loads(request.form.get("reportData", "[]"))

        doc = Document()
        title = doc.add_heading("스마트 제조 보고서", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        author = doc.add_paragraph("작성자 :        ")
        author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for rep in report_data:
            doc.add_heading(f"{rep['process']} 공정 보고서", level=1)
            doc.add_paragraph(f"분석 기간: {rep.get('range', '')}")

            if "summary" in rep:
                doc.add_paragraph("요약", style="Heading 1")
                doc.add_paragraph(rep["summary"])

            if "available" in rep:
                avg = round(sum(rep["available"]) / len(rep["available"]) * 100, 1)
                doc.add_paragraph(f"평균 가동률: {avg}%")
                donut = create_donut_chart(avg, ["가동률", "비가동률"], ["green", "#e0e0e0"])
                doc.add_picture(donut, width=Inches(2.5))

            if "production" in rep:
                p0 = rep["production"]["input"]
                p3 = rep["production"]["output"]
                rate = rep["production"]["rate"]
                doc.add_paragraph(f"생산실적: 투입 {p0} → 산출 {p3} (양품률 {rate}%)")
                bar = create_production_bar(p0, p3)
                doc.add_picture(bar, width=Inches(3))

            if "mtbf_minutes" in rep and "total_processing_minutes" in rep and "failure_count" in rep:
                doc.add_paragraph("MTBF 요약")
                doc.add_paragraph("---------------------------")
                doc.add_paragraph(f"고장 횟수: {rep['failure_count']}회")
                doc.add_paragraph(f"총 운영 시간: {rep['total_processing_minutes']}분")
                doc.add_paragraph(f"고장 간 평균 시간 (MTBF): {rep['mtbf_minutes']}분")

            if "mttr_minutes" in rep and "total_repair_minutes" in rep and "repair_count" in rep:
                doc.add_paragraph("MTTR 요약")
                doc.add_paragraph("---------------------------")
                doc.add_paragraph(f"수리 횟수: {rep['repair_count']}회")
                doc.add_paragraph(f"총 수리 시간: {rep['total_repair_minutes']}분")
                doc.add_paragraph(f"평균 수리 시간 (MTTR): {rep['mttr_minutes']}분")

            if "failureLabels" in rep and "failureCounts" in rep:
                doc.add_paragraph("고장 발생 분포")
                line = create_failure_line_chart(rep["failureLabels"], rep["failureCounts"])
                doc.add_picture(line, width=Inches(4))

            if "failure_total" in rep and "maintenance_total" in rep and "total_processing_minutes" in rep:
                doc.add_paragraph("총 다운타임 분석")
                pie = create_total_downtime_pie(
                    rep["failure_total"], rep["maintenance_total"], rep["total_processing_minutes"]
                )
                doc.add_picture(pie, width=Inches(3))

            if "downtime_hour_labels" in rep and "failure_by_hour" in rep:
                doc.add_paragraph("시간대별 고장 다운타임")
                bar = create_downtime_bar_chart(rep["downtime_hour_labels"], rep["failure_by_hour"])
                doc.add_picture(bar, width=Inches(4.5))

            if "maintenance_total" in rep:
                doc.add_paragraph("정비 요약")
                doc.add_paragraph(f"총 정비 시간: {rep['maintenance_total']}분")

        # 부록: 고장 분포 테이블
        doc.add_page_break()
        doc.add_heading("전체 고장 발생 분포 테이블", level=1)
        hour_map = {}
        for label, count in zip(failure_labels, failure_counts):
            hour = label[:2] + "시대"
            hour_map[hour] = hour_map.get(hour, 0) + count

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "시간대"
        hdr[1].text = "고장 수"
        for hour, cnt in sorted(hour_map.items()):
            row = table.add_row().cells
            row[0].text = hour
            row[1].text = str(cnt)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")

    except Exception as e:
        print("📄 DOCX 생성 오류:", e)
        return jsonify({"error": "파일 생성 실패"}), 500


@app.route("/generate_excel", methods=["POST"])
def generate_excel():
    try:
        from openpyxl import Workbook
        from datetime import timezone, timedelta
        from influxdb_client import InfluxDBClient
        import os

        report_data = json.loads(request.form.get("reportData", "[]"))
        wb = Workbook()

        INFLUX_URL = os.getenv("INFLUX_URL")
        INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
        INFLUX_ORG = os.getenv("INFLUX_ORG")
        influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)

        all_logs = []
        downtime_states = ["failure", "repair"]

        for rep in report_data:
            process = rep["process"]
            bucket_name = f"{process}_status"
            range_str = rep.get("range", "1h")

            if "/" in range_str:
                start, end = range_str.split("/")
                range_clause = f'|> range(start: time(v: "{start}"), stop: time(v: "{end}"))'
            else:
                range_clause = f'|> range(start: -{range_str})'

            query = f'''
            from(bucket: "{bucket_name}")
              {range_clause}
              |> filter(fn: (r) => r._measurement == "status_log")
              |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
              |> keep(columns: ["_time", "available", "event_type"])
            '''
            result = influx_client.query_api().query(query)

            for table in result:
                for record in table.records:
                    time_obj = record.get_time().astimezone(timezone(timedelta(hours=9)))
                    event_type = record.values.get("event_type", "")
                    is_downtime = "O" if event_type in downtime_states else "X"

                    all_logs.append([
                        time_obj.strftime("%Y-%m-%d %H:%M:%S"),
                        record.values.get("available", ""),
                        event_type,
                        process,
                        is_downtime,
                        time_obj.strftime("%H시대")
                    ])

            # 생산 실적 시트 작성
            prod_ws = wb["생산 실적"] if "생산 실적" in wb.sheetnames else wb.create_sheet(title="생산 실적")
            if prod_ws.max_row == 1:
                prod_ws.append(["시간", "공정 ID", "제품 ID"])

            for proc_id in ["P0", "P3"]:
                prod_query = f'''
                from(bucket: "process")
                  {range_clause}
                  |> filter(fn: (r) => r._measurement == "process_log" and r.process_id == "{proc_id}")
                  |> keep(columns: ["_time", "process_id", "product_id"])
                '''
                result = influx_client.query_api().query(org=INFLUX_ORG, query=prod_query)

                for table in result:
                    for record in table.records:
                        time_str = record.get_time().astimezone(timezone(timedelta(hours=9))).strftime("%Y-%m-%d %H:%M:%S")
                        prod_ws.append([time_str, record.values["process_id"], record.values["product_id"]])

        # 공정 이력 시트
        log_sheet = wb.active
        log_sheet.title = "공정 이력"
        log_sheet.append(["시간", "가동여부", "이벤트 타입", "공정", "다운타임 여부", "시대"])
        for row in sorted(all_logs, key=lambda r: r[0], reverse=True):
            log_sheet.append(row)

        # 생산 실적 시트 정렬
        prod_data = list(prod_ws.iter_rows(min_row=2, values_only=True))
        prod_data_sorted = sorted(prod_data, key=lambda x: x[0], reverse=True)
        for i in range(prod_ws.max_row, 1, -1):
            prod_ws.delete_rows(i)
        for row in prod_data_sorted:
            prod_ws.append(row)

        output = BytesIO()
        wb.save(output)
        output.seek(0)

        response = make_response(output.read())
        response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response.headers.set('Content-Disposition', 'attachment; filename=제조_기초데이터.xlsx')
        return response

    except Exception as e:
        print("❌ Excel 생성 오류:", e)
        return jsonify({"error": "엑셀 파일 다운로드 실패"}), 500

# ===============================
# 서버 실행
# ===============================
if __name__ == "__main__":
    socketio.start_background_task(target=emit_status)
    socketio.run(app, host='0.0.0.0', port=5000, debug=True, use_reloader=False)
