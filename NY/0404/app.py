import eventlet
eventlet.monkey_patch()

import os
from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO
from influxdb_client import InfluxDBClient
from dotenv import load_dotenv
from flask_cors import CORS
from docx import Document
from io import BytesIO
from docx.shared import Inches
import base64
import openai
from collections import defaultdict
from datetime import datetime, timezone, timedelta
import json
import traceback

# ===============================
# 환경 변수 및 클라이언트 설정
# ===============================
load_dotenv()

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*", async_mode="eventlet")

INFLUX_URL = os.getenv("INFLUX_URL")
INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
INFLUX_ORG = os.getenv("INFLUX_ORG")
influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)
openai.api_key = os.getenv("OPENAI_API_KEY")

# ===============================
# 한글 기간 문자열 변환 함수
# ===============================
def normalize_range(range_str):
    kor_to_influx = {
        "1시간": "1h", "3시간": "3h", "6시간": "6h", "9시간": "9h",
        "1일": "1d", "7일": "7d", "31일": "31d"
    }
    return kor_to_influx.get(range_str, range_str)

# ===============================
# 실시간 상태 조회 및 전송
# ===============================
def get_recent_status(bucket):
    query = f'''
    from(bucket: "{bucket}")
      |> range(start: -30s)
      |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type")
      |> sort(columns: ["_time"], desc: true)
      |> limit(n: 3)
    '''
    result = influx_client.query_api().query(org=INFLUX_ORG, query=query)
    events = [record.get_value() for table in result for record in table.records]
    return events if events else None

def emit_status():
    prev_events = {"P1-A": None, "P1-B": None, "P2-A": None, "P2-B": None}
    while True:
        for key in prev_events.keys():
            events = get_recent_status(f"{key}_status")
            if events:
                latest = events[0]
                if latest != prev_events[key]:
                    socketio.emit('status_update', {key: {'event_type': latest}})
                    prev_events[key] = latest
        socketio.sleep(1)

@socketio.on('connect')
def handle_connect():
    for key in ["P1", "P2"]:
        events = get_recent_status(f"{key}_status")
        if events:
            latest = events[0]
            socketio.emit('status_update', {f"{key}-A": {'event_type': latest}})

# ===============================
# 라우팅
# ===============================
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/report")
def report_page():
    return render_template("report.html")

# ===============================
# 보고서 생성 API (다중 공정 대응)
# ===============================
@app.route("/generate_report", methods=["POST"])
def generate_report():
    try:
        data = request.get_json()
        processes = data.get("processes", [])
        range_str = normalize_range(data.get("range"))

        all_reports = []
        for process in processes:
            if "/" in range_str:
                start, end = range_str.split("/")
                range_clause = f'|> range(start: time(v: "{start}"), stop: time(v: "{end}"))'
                start_time = datetime.fromisoformat(start.replace("+09:00", ""))
                end_time = datetime.fromisoformat(end.replace("+09:00", ""))
            else:
                range_clause = f'|> range(start: -{range_str})'
                start_time, end_time = None, None

            query = f'''
            from(bucket: "{process}_status")
              {range_clause}
              |> filter(fn: (r) => r._measurement == "status_log")
              |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
              |> keep(columns: ["_time", "available", "event_type"])
            '''
            tables = influx_client.query_api().query(query)

            total, available_sum, failure_count = 0, 0, 0
            time_labels, available_values, failure_values = [], [], []
            failure_hourly = defaultdict(int)

            KST = timezone(timedelta(hours=9))

            for table in tables:
                for record in table.records:
                    event_type = record.get_value()
                    timestamp = record.get_time().astimezone(KST).replace(tzinfo=None)

                    # 1️⃣ 기간 필터
                    if start and end:
                        if not (start <= timestamp <= end):
                            continue

                    # 2️⃣ 운영 시간 필터 (09시~18시)
                    if not (9 <= timestamp.hour < 18):
                        continue

                    # 3️⃣ 이벤트 처리
                    if event_type in ["failure", "maintenance"]:
                        current_event = event_type
                        current_start = timestamp

                    elif event_type == "processing" and current_event and current_start:
                        diff = (timestamp - current_start).total_seconds() / 60  # 분 단위
                        hour_label = current_start.strftime("%H시")

                        if current_event == "failure":
                            failure_total += diff
                            failure_by_hour[hour_label] += diff
                        elif current_event == "maintenance":
                            maintenance_total += diff
                            maintenance_by_hour[hour_label] += diff

                        current_event = None
                        current_start = None

                    available = record.values.get("available", 0)
                    event_type = record.values.get("event_type", "")
                    timestamp = record_time.strftime("%H:%M")
                    hour_label = record_time.strftime("%H시대")

                    total += 1
                    available_sum += available
                    if event_type == "failure":
                        failure_count += 1
                        failure_hourly[hour_label] += 1

                    time_labels.append(timestamp)
                    available_values.append(round(available, 2))
                    failure_values.append(1 if event_type == "failure" else 0)

            avg_avail = round((available_sum / total) * 100, 1) if total else 0
            failure_table_labels = list(failure_hourly.keys())
            failure_table_counts = list(failure_hourly.values())

            prompt = f"""
            공정명: {process}
            기간: 최근 {range_str}
            가동률 평균: {avg_avail}%
            고장 횟수: {failure_count}회

            위 데이터를 바탕으로 제조 공정 보고서를 작성해줘. 다음 항목을 포함해줘:
            1. 공정 요약
            2. 주요 이슈
            3. 대응 조치
            4. 향후 제언
            """
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                    {"role": "user", "content": prompt}
                ]
            )
            all_reports.append({
                "process": process,
                "report": response.choices[0].message.content,
                "labels": time_labels,
                "available": available_values,
                "failures": failure_values,
                "failureLabels": failure_table_labels,
                "failureCounts": failure_table_counts
            })

        return jsonify({"reports": all_reports})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ===============================
# DOCX 다운로드 API (시간대 그룹핑 반영)
# ===============================
from flask import request

@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        text = request.form.get("report", "")
        failure_labels = request.form.get("failureLabels", "[]")
        failure_counts = request.form.get("failureCounts", "[]")

        doc = Document()
        doc.add_heading("📄 스마트 제조 보고서", 0)
        doc.add_paragraph(text)

        # ✅ 이미지 추가
        avail_imgs = request.files.getlist("availabilityImages")
        fail_imgs = request.files.getlist("failureImages")

        for i in range(len(avail_imgs)):
            doc.add_paragraph(f"✅ [{i+1}] 가동률 변화")
            doc.add_picture(BytesIO(avail_imgs[i].read()), width=Inches(4))

        for i in range(len(fail_imgs)):
            doc.add_paragraph(f"📊 [{i+1}] 고장 발생 분포")
            doc.add_picture(BytesIO(fail_imgs[i].read()), width=Inches(4))

        # ✅ 고장 테이블
        import json
        labels = json.loads(failure_labels)
        counts = json.loads(failure_counts)

        hour_map = {}
        for label, count in zip(labels, counts):
            hour = label[:2] + "시대"
            hour_map[hour] = hour_map.get(hour, 0) + count

        doc.add_paragraph("📊 고장 발생 분포 테이블 (시간대 기준)")
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '시간대'
        hdr_cells[1].text = '고장 수'
        for hour, count in hour_map.items():
            row_cells = table.add_row().cells
            row_cells[0].text = hour
            row_cells[1].text = str(count)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")
    except Exception as e:
        print(f"📄 DOCX 생성 오류: {e}")
        return jsonify({"error": "파일 생성 실패"}), 500


# ===============================
# 다운타임 계산 API
# ===============================
@app.route("/get_downtime_data", methods=["POST"])
def get_downtime_data():
    from datetime import datetime, timezone, timedelta
    from collections import defaultdict

    data = request.json
    process = data.get("process")
    range_str = data.get("range")

    if not process or not range_str:
        return jsonify({"error": "Missing process or range"}), 400

    # 시간 범위 파싱
    if "/" in range_str:
        start_str, end_str = range_str.split("/")
        start = datetime.fromisoformat(start_str.replace("Z", "").replace("+09:00", ""))
        end = datetime.fromisoformat(end_str.replace("Z", "").replace("+09:00", ""))
        range_clause = f'|> range(start: time(v: "{start_str}"), stop: time(v: "{end_str}"))'
    else:
        range_clause = f'|> range(start: -{range_str})'
        now = datetime.utcnow()
        if "h" in range_str:
            hours = int(range_str.replace("h", ""))
            start = now - timedelta(hours=hours)
        elif "d" in range_str:
            days = int(range_str.replace("d", ""))
            start = now - timedelta(days=days)
        end = now

    query = f'''
    from(bucket: "{process}_status")
      {range_clause}
      |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type")
      |> sort(columns: ["_time"])
    '''

    tables = influx_client.query_api().query(org=INFLUX_ORG, query=query)

    # 다운타임 누적
    failure_total = 0
    maintenance_total = 0
    failure_by_hour = defaultdict(int)
    maintenance_by_hour = defaultdict(int)

    current_event = None
    current_start = None

    KST = timezone(timedelta(hours=9))

    for table in tables:
        for record in table.records:
            event_type = record.get_value()
            timestamp = record.get_time().astimezone(KST).replace(tzinfo=None)

            if event_type in ["failure", "maintenance"]:
                current_event = event_type
                current_start = timestamp
            elif event_type == "processing" and current_event and current_start:
                diff = (timestamp - current_start).total_seconds() / 60  # 분 단위
                hour_label = current_start.strftime("%H시")

                if current_event == "failure":
                    failure_total += diff
                    failure_by_hour[hour_label] += diff
                elif current_event == "maintenance":
                    maintenance_total += diff
                    maintenance_by_hour[hour_label] += diff

                current_event = None
                current_start = None

    return jsonify({
        "failure_total": round(failure_total, 1),
        "maintenance_total": round(maintenance_total, 1),
        "hourly_labels": sorted(set(list(failure_by_hour.keys()) + list(maintenance_by_hour.keys()))),
        "failure_by_hour": [round(failure_by_hour[h], 1) for h in sorted(failure_by_hour.keys())],
        "maintenance_by_hour": [round(maintenance_by_hour[h], 1) for h in sorted(maintenance_by_hour.keys())]
    })


# ===============================
# 가동률 계산 API
# ===============================
@app.route("/calculate_availability", methods=["POST"])
def calculate_availability():
    try:
        data = request.json
        process = data.get("process")  # 예시: "P1-A"
        period = data.get("period")    # "일간", "주간", "월간"

        # 기간 설정
        now = datetime.utcnow()
        if period == "일간":
            start = now - timedelta(hours=24)
        elif period == "주간":
            start = now - timedelta(days=7)
        elif period == "월간":
            start = now - timedelta(days=31)
        else:
            return jsonify({"error": "Invalid period"}), 400

        start_str = start.isoformat() + "Z"
        end_str = now.isoformat() + "Z"

        # 쿼리문
        query = f'''
        from(bucket: "{process}_status")
          |> range(start: {start_str}, stop: {end_str})
          |> filter(fn: (r) => r._measurement == "status_log" and r._field == "available")
        '''

        tables = influx_client.query_api().query(org=INFLUX_ORG, query=query)

        # 데이터 수집 및 계산
        total, available_sum = 0, 0
        for table in tables:
            for record in table.records:
                timestamp = record.get_time().astimezone(KST).replace(tzinfo=None)
                hour = timestamp.hour

                # ✅ 운영 시간 필터: 09시 ~ 18시만 포함
                if not (9 <= hour < 18):
                    continue

                value = record.get_value()
                if value is not None:
                    total += 1
                    available_sum += value

        availability = round((available_sum / total) * 100, 1) if total else 0

        return jsonify({
            "process": process,
            "period": period,
            "availability": availability
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ===============================
# 고장건수 API 추가
# ===============================
@app.route("/calculate_failure_count", methods=["POST"])
def calculate_failure_count():
    try:
        data = request.json
        process = data.get("process")  # 예: "P1-A"
        period = data.get("period")    # "일간", "주간", "월간"

        # 기간 설정
        now = datetime.utcnow()
        if period == "일간":
            start = now - timedelta(hours=24)
        elif period == "주간":
            start = now - timedelta(days=7)
        elif period == "월간":
            start = now - timedelta(days=31)
        else:
            return jsonify({"error": "Invalid period"}), 400

        start_str = start.isoformat() + "Z"
        end_str = now.isoformat() + "Z"

        # 쿼리문
        query = f'''
        from(bucket: "{process}_status")
        |> range(start: {start_str}, stop: {end_str})
        |> filter(fn: (r) => r._measurement == "status_log" and r._field == "event_type" and r._value == "failure")
        '''

        tables = influx_client.query_api().query(org=INFLUX_ORG, query=query)

        KST = timezone(timedelta(hours=9))
        failure_count = 0

        for table in tables:
            for record in table.records:
                timestamp = record.get_time().astimezone(KST).replace(tzinfo=None)
                hour = timestamp.hour

                # ✅ 운영 시간만 포함
                if not (9 <= hour < 18):
                    continue

                failure_count += 1

        return jsonify({
            "process": process,
            "period": period,
            "failure_count": failure_count
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===============================
# 서버 실행
# ===============================
if __name__ == "__main__":
    socketio.start_background_task(target=emit_status)
    socketio.run(app, host='0.0.0.0', port=5000, debug=False, use_reloader=False)

