from flask import Flask, request, jsonify, render_template, send_file
from influxdb_client import InfluxDBClient
from docx import Document
from docx.shared import Inches
from io import BytesIO
from openai import OpenAI  # ✅ 올바른 방식
import base64
import os
from dotenv import load_dotenv

# ✅ 환경 변수 로드
load_dotenv()

# ✅ Flask 앱 초기화
app = Flask(__name__)

# ✅ .env에서 민감 정보 불러오기
openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
INFLUX_URL = os.getenv("INFLUX_URL")
INFLUX_TOKEN = os.getenv("INFLUX_TOKEN")
INFLUX_ORG = os.getenv("INFLUX_ORG")

influx_client = InfluxDBClient(url=INFLUX_URL, token=INFLUX_TOKEN, org=INFLUX_ORG)

# ✅ 메인 대시보드 페이지
@app.route("/")
def index():
    return render_template("index.html")

# ✅ 보고서 페이지
@app.route("/report")
def report_page():
    return render_template("report.html")

# ✅ 보고서 생성 API
@app.route("/generate_report", methods=["POST"])
def generate_report():
    data = request.json
    process = data.get("process")
    range_str = data.get("range")

    query = f'''
    from(bucket: "{process}_status")
      |> range(start: -{range_str})
      |> filter(fn: (r) => r._measurement == "status_log")
      |> pivot(rowKey:["_time"], columnKey: ["_field"], valueColumn: "_value")
      |> keep(columns: ["_time", "available", "event_type"])
    '''

    tables = influx_client.query_api().query(query)

    total = 0
    available_sum = 0
    failure_count = 0
    time_labels = []
    available_values = []
    failure_values = []

    for table in tables:
        for record in table.records:
            available = record.values.get("available", 0)
            event_type = record.values.get("event_type", "")
            timestamp = record.values["_time"].strftime("%H:%M")

            total += 1
            available_sum += available
            if event_type == "failure":
                failure_count += 1

            time_labels.append(timestamp)
            available_values.append(round(available, 2))
            failure_values.append(1 if event_type == "failure" else 0)

    avg_avail = round((available_sum / total) * 100, 1) if total else 0

    prompt = f"""
공정명: {process}
기간: 최근 {range_str}
가동률 평균: {avg_avail}%
고장 횟수: {failure_count}회

위 데이터를 바탕으로 제조 공정 보고서를 작성해줘. 다음 항목을 포함해줘:
1. 공정 요약
2. 주요 이슈
3. 대응 조치
4. 향후 제언
"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "너는 제조공정 보고서를 작성하는 AI 비서야."},
                {"role": "user", "content": prompt}
            ]
        )
        return jsonify({
            "report": response.choices[0].message.content,
            "labels": time_labels,
            "available": available_values,
            "failures": failure_values
        })
    except Exception as e:
        print(f"오류 발생: {e}")
        return jsonify({"error": str(e)}), 500

# ✅ 보고서 다운로드 API
@app.route("/generate_docx", methods=["POST"])
def generate_docx():
    try:
        data = request.json
        text = data.get("report", "")
        avail_img_b64 = data.get("availabilityImage", "")
        fail_img_b64 = data.get("failureImage", "")

        doc = Document()
        doc.add_heading("📄 스마트 제조 보고서", 0)
        doc.add_paragraph(text)

        def add_image(doc, b64_string, title):
            if b64_string and "base64," in b64_string:
                try:
                    doc.add_paragraph(title)
                    image_data = base64.b64decode(b64_string.split(",")[-1])
                    image_stream = BytesIO(image_data)
                    doc.add_picture(image_stream, width=Inches(2.75))
                except Exception as img_err:
                    print(f"이미지 디코딩 오류: {img_err}")

        add_image(doc, avail_img_b64, "✅ 가동률 변화")
        add_image(doc, fail_img_b64, "📊 고장 발생 분포")

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name="제조_보고서.docx")

    except Exception as e:
        print(f"📛 DOCX 생성 오류: {e}")
        return jsonify({"error": "파일 생성 실패"}), 500

# ✅ 챗봇 응답 처리 API
@app.route("/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message", "")

    messages = [
        {"role": "system", "content": "너는 제조 상태와 데이터를 알려주는 AI Agent야."},
        {"role": "user", "content": user_message}
    ]

    try:
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=messages
        )
        return jsonify({"reply": response.choices[0].message.content})
    except Exception as e:
        return jsonify({"reply": f"⚠️ 오류 발생: {str(e)}"}), 500

# ✅ 앱 실행
if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
